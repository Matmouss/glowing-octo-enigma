import os
import re
from docx import Document

def normalize_text(text):
    return re.sub(r'[^a-zA-Z]', '', text.lower())

def parse_txt_elements(txt_path):
    with open(txt_path, 'r', encoding='utf-8') as file:
        content = file.read()
    elements = [element.strip() for element in content.split(';') if element.strip()]
    return elements

def remove_elements_from_docx(doc_path, output_path, elements_to_remove):
    doc = Document(doc_path)
    normalized_elements = [normalize_text(element) for element in elements_to_remove]
    words_removed_count = 0

    for paragraph in doc.paragraphs:
        words = paragraph.text.split()
        new_words = [word for word in words if normalize_text(word) not in normalized_elements]
        words_removed_count += len(words) - len(new_words)
        paragraph.text = ' '.join(new_words)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                words = cell.text.split()
                new_words = [word for word in words if normalize_text(word) not in normalized_elements]
                words_removed_count += len(words) - len(new_words)
                cell.text = ' '.join(new_words)

    doc.save(output_path)
    return words_removed_count

current_dir = os.path.dirname(os.path.abspath(__file__))
doc_path = os.path.join(current_dir, "fichier_html.docx")
output_path = os.path.join(current_dir, "document_modifie.docx")
txt_path = os.path.join(current_dir, "elements.txt")
elements_to_remove = parse_txt_elements(txt_path)

words_removed = remove_elements_from_docx(doc_path, output_path, elements_to_remove)
print("Traitement terminé. Document modifié enregistré sous", output_path)
print("Nombre total de mots supprimés :", words_removed)
